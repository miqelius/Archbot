"""
services/document_processor.py - Extract text from various file formats
Supports PDF, DOCX, images (with OCR)
"""

import asyncio
import logging
from pathlib import Path
from typing import Optional
from io import BytesIO

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Unified interface for extracting text from documents.
    Handles PDF, DOCX, images with automatic format detection.
    """
    
    def __init__(self):
        self.logger = logging.getLogger("DocumentProcessor")
    
    async def extract_text(
        self,
        file_path: str,
    ) -> str:
        """
        Extract text from document.
        Auto-detects format from file extension or magic bytes.
        
        Args:
            file_path: Path to file or file URL
        
        Returns:
            Extracted text
        
        Raises:
            ValueError: If format not supported
            IOError: If file cannot be read
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        self.logger.info(f"Extracting text from {file_path} ({extension})")
        
        try:
            if extension == ".pdf":
                text = await self._extract_pdf(str(path))
            elif extension in {".docx", ".doc"}:
                text = await self._extract_docx(str(path))
            elif extension in {".jpg", ".jpeg", ".png", ".gif"}:
                text = await self._extract_image_ocr(str(path))
            else:
                raise ValueError(f"Unsupported file format: {extension}")
            
            self.logger.info(f"Extracted {len(text)} characters")
            return text
        
        except Exception as e:
            self.logger.error(f"Extraction failed: {e}")
            raise
    
    async def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        try:
            import PyPDF2
            
            text_parts = []
            
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                            self.logger.debug(
                                f"Extracted page {page_num + 1}: "
                                f"{len(page_text)} chars"
                            )
                    except Exception as e:
                        self.logger.warning(f"Failed to extract page {page_num}: {e}")
            
            full_text = "\n".join(text_parts)
            
            if not full_text.strip():
                # Try OCR for image-based PDF
                self.logger.info("PDF appears image-based, attempting OCR...")
                full_text = await self._extract_pdf_ocr(file_path)
            
            return full_text
        
        except ImportError:
            raise ImportError("PyPDF2 not installed. Install with: pip install PyPDF2")
    
    async def _extract_pdf_ocr(self, file_path: str) -> str:
        """Extract text from PDF using OCR (fallback for image-based PDFs)"""
        try:
            import pdf2image
            import pytesseract
            from PIL import Image
            
            text_parts = []
            
            # Convert PDF to images
            images = pdf2image.convert_from_path(file_path)
            
            for page_num, image in enumerate(images):
                try:
                    # OCR on image
                    page_text = pytesseract.image_to_string(image)
                    if page_text:
                        text_parts.append(page_text)
                        self.logger.debug(
                            f"OCR page {page_num + 1}: {len(page_text)} chars"
                        )
                except Exception as e:
                    self.logger.warning(f"OCR failed for page {page_num}: {e}")
            
            return "\n".join(text_parts)
        
        except ImportError:
            self.logger.error(
                "OCR libraries not installed. Install with: "
                "pip install pdf2image pytesseract pillow"
            )
            raise
    
    async def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX/DOC files"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
        
        except ImportError:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
    
    async def _extract_image_ocr(self, file_path: str) -> str:
        """Extract text from image using OCR"""
        try:
            import pytesseract
            from PIL import Image
            
            # Open image
            image = Image.open(file_path)
            
            # OCR
            text = pytesseract.image_to_string(image)
            
            return text
        
        except ImportError:
            raise ImportError(
                "OCR libraries not installed. Install with: "
                "pip install pytesseract pillow"
            )
    
    async def extract_batch(
        self,
        file_paths: list[str],
    ) -> dict[str, str]:
        """
        Extract text from multiple documents concurrently.
        
        Args:
            file_paths: List of file paths
        
        Returns:
            Mapping of file_path → extracted_text
        """
        tasks = [self.extract_text(path) for path in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        output = {}
        for file_path, result in zip(file_paths, results):
            if isinstance(result, Exception):
                self.logger.error(f"Failed to extract {file_path}: {result}")
                output[file_path] = f"[Error: {str(result)}]"
            else:
                output[file_path] = result
        
        return output
    
    async def extract_from_bytes(
        self,
        file_bytes: bytes,
        file_name: str,
    ) -> str:
        """
        Extract text from file bytes (for uploaded files).
        
        Args:
            file_bytes: Raw file bytes
            file_name: Original filename (for format detection)
        
        Returns:
            Extracted text
        """
        # Write to temp file
        import tempfile
        
        with tempfile.NamedTemporaryFile(
            suffix=Path(file_name).suffix,
            delete=False
        ) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        
        try:
            text = await self.extract_text(tmp_path)
            return text
        finally:
            # Clean up temp file
            Path(tmp_path).unlink(missing_ok=True)


# Convenience function
processor = DocumentProcessor()
